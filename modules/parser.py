import PyPDF2
import pdfplumber
from docx import Document
import re

class ContractParser:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_document(self, file_path, file_type):
        """Parse uploaded contract document"""
        if file_type == 'pdf':
            return self._parse_pdf(file_path)
        elif file_type == 'docx':
            return self._parse_docx(file_path)
        else:
            return self._parse_txt(file_path)
    
    def _parse_pdf(self, file_path):
        """Extract text from PDF"""
        text = ""
        metadata = {}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except Exception as e:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
            except Exception as e2:
                text = "Could not extract text from PDF"
                metadata['error'] = str(e2)
        
        if not text.strip():
            text = "No text extracted from PDF"
        
        sections = self._detect_sections(text)
        
        return {
            'raw_text': text,
            'sections': sections,
            'metadata': metadata
        }
    
    def _parse_docx(self, file_path):
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            sections = self._detect_sections(text)
            return {
                'raw_text': text,
                'sections': sections,
                'metadata': {'paragraphs': len(doc.paragraphs)}
            }
        except Exception as e:
            return {
                'raw_text': "Error reading Word document",
                'sections': [],
                'metadata': {'error': str(e)}
            }
    
    def _parse_txt(self, file_path):
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
            except Exception as e:
                return {
                    'raw_text': "Error reading text file",
                    'sections': [],
                    'metadata': {'error': str(e)}
                }
        
        sections = self._detect_sections(text)
        return {
            'raw_text': text,
            'sections': sections,
            'metadata': {}
        }
    
    def _detect_sections(self, text):
        """Detect contract sections"""
        if not text or len(text) < 10:
            return []
        
        sections = []
        patterns = [
            r'(?i)(article|section|clause)\s+(\d+)',
            r'(?i)(parties|definitions|scope|term|payment|liability)'
        ]
        
        for pattern in patterns:
            try:
                matches = re.finditer(pattern, text)
                for match in matches:
                    sections.append({
                        'header': match.group(0),
                        'position': match.start()
                    })
            except:
                pass
        
        return sections